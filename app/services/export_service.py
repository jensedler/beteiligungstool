import os
import tempfile

import markdown
from docx import Document
from docx.shared import Pt

from app.models.konzept import Konzept


def _get_text(konzept: Konzept) -> str:
    return konzept.edited_text or konzept.generated_text or ""


def export_markdown(konzept: Konzept) -> str:
    text = f"# {konzept.title}\n\n{_get_text(konzept)}"
    fd, path = tempfile.mkstemp(suffix=".md")
    with os.fdopen(fd, "w", encoding="utf-8") as f:
        f.write(text)
    return path


def export_docx(konzept: Konzept) -> str:
    doc = Document()
    doc.add_heading(konzept.title, level=0)

    text = _get_text(konzept)
    for line in text.split("\n"):
        line = line.rstrip()
        if line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("- ") or line.startswith("* "):
            doc.add_paragraph(line[2:], style="List Bullet")
        elif line.strip():
            doc.add_paragraph(line)

    fd, path = tempfile.mkstemp(suffix=".docx")
    os.close(fd)
    doc.save(path)
    return path
